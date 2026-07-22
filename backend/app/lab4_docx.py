"""Lab 4 — render a structured, production-grade job aid into a branded .docx.

Opens a template canvas (see lab4_templates) and appends a full government
document: control block, roles, a procedure with decision branches and callout
boxes, a quick-reference card, definitions, revision history, and a sign-off.
"""
import io
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# label, text colour, fill colour
_CALLOUT = {
    "warning": ("WARNING", "C0392B", "FBE7E7"),
    "caution": ("CAUTION", "B7791F", "FBEED6"),
    "note": ("NOTE", "2F49B5", "E8ECFD"),
}
_HAIRLINE = "D0D9E8"


def _para(doc, text="", *, italic=False, bold=False, size=None, color=None, indent=None):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.italic = italic
        r.bold = bold
        if size is not None:
            r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor.from_string(color)
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    return p


def _shade(cell, fill: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shd)


def _borders(table, color: str) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _cell_text(cell, text, *, bold=False, size=10, color=None) -> None:
    r = cell.paragraphs[0].add_run("" if text is None else str(text))
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def _kv_table(doc, pairs) -> None:
    """Two-column field/value table (document control)."""
    table = doc.add_table(rows=0, cols=2)
    for label, value in pairs:
        if not value:
            continue
        cells = table.add_row().cells
        _shade(cells[0], "EEF2F9")
        _cell_text(cells[0], label, bold=True, size=9)
        _cell_text(cells[1], value)
    _borders(table, _HAIRLINE)


def _grid_table(doc, headers, rows, brand) -> None:
    """Header-row table (roles, definitions, revision history, approvals)."""
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        _shade(table.rows[0].cells[i], brand)
        _cell_text(table.rows[0].cells[i], h, bold=True, size=9, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _cell_text(cells[i], val)
    _borders(table, _HAIRLINE)


def _callout(doc, kind, text) -> None:
    label, color, fill = _CALLOUT.get(kind, _CALLOUT["note"])
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(color)
    p.add_run(str(text)).font.size = Pt(10)
    _borders(table, color)


def _heading(doc, text) -> None:
    doc.add_paragraph(text, style="Heading 1")


def render_job_aid(job_aid: dict, base_path: Path, brand: str = "1F3A93") -> bytes:
    doc = Document(str(base_path))
    control = job_aid.get("control") or {}

    # ── Title + one-line meta ──
    doc.add_paragraph(job_aid.get("title") or "Job Aid", style="Title")
    meta = [job_aid.get("document_type") or "Job Aid"]
    if job_aid.get("audience"):
        meta.append(job_aid["audience"])
    if control.get("classification"):
        meta.append(control["classification"])
    _para(doc, "  ·  ".join(meta), size=9, color="6B7280")

    # ── Document control ──
    _heading(doc, "Document Control")
    _kv_table(doc, [
        ("Document ID", control.get("document_id")),
        ("Version", control.get("version")),
        ("Effective date", control.get("effective_date")),
        ("Next review", control.get("review_date")),
        ("Owner", control.get("owner")),
        ("Approver", control.get("approver")),
        ("Classification", control.get("classification")),
    ])

    # ── Purpose & scope ──
    if job_aid.get("purpose"):
        _heading(doc, "Purpose")
        _para(doc, job_aid["purpose"])
    if job_aid.get("scope"):
        _heading(doc, "Scope")
        _para(doc, job_aid["scope"])

    # ── Roles & responsibilities ──
    roles = job_aid.get("roles") or []
    if roles:
        _heading(doc, "Roles & Responsibilities")
        _grid_table(doc, ["Role", "Responsibility"],
                    [[r.get("role", ""), r.get("responsibility", "")] for r in roles],
                    brand)

    # ── Prerequisites ──
    prerequisites = job_aid.get("prerequisites") or []
    if prerequisites:
        _heading(doc, "Before You Start")
        for item in prerequisites:
            doc.add_paragraph(str(item), style="List Bullet")

    # ── Procedure ──
    for section in job_aid.get("procedure") or []:
        _heading(doc, section.get("heading") or "Procedure")
        for i, step in enumerate(section.get("steps") or [], 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {step.get('title', '')}").bold = True
            if step.get("role"):
                r = p.add_run(f"   [{step['role']}]")
                r.italic = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor.from_string("6B7280")
            detail = step.get("detail")
            if detail and detail != step.get("title"):
                _para(doc, detail, indent=0.3)
            decision = step.get("decision")
            if decision:
                _para(doc, decision.get("question", "Decision:"), bold=True, indent=0.3)
                for br in decision.get("branches") or []:
                    _para(doc, f"{br.get('condition', '')} → {br.get('action', '')}", indent=0.55)
            callout = step.get("callout")
            if callout:
                _callout(doc, callout.get("type", "note"), callout.get("text", ""))

    # ── Quick reference card ──
    quick = job_aid.get("quick_reference") or []
    if quick:
        _heading(doc, "Quick Reference Card")
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        _shade(cell, "F4F7FC")
        for i, item in enumerate(quick, 1):
            para = cell.paragraphs[0] if i == 1 else cell.add_paragraph()
            para.add_run(f"{i}. {item}").font.size = Pt(10)
        _borders(table, _HAIRLINE)

    # ── Definitions ──
    definitions = job_aid.get("definitions") or []
    if definitions:
        _heading(doc, "Definitions & Acronyms")
        _grid_table(doc, ["Term", "Definition"],
                    [[d.get("term", ""), d.get("definition", "")] for d in definitions],
                    brand)

    # ── Revision history ──
    revisions = job_aid.get("revision_history") or []
    if revisions:
        _heading(doc, "Revision History")
        _grid_table(doc, ["Version", "Date", "Author", "Summary"],
                    [[r.get("version", ""), r.get("date", ""), r.get("author", ""), r.get("summary", "")]
                     for r in revisions],
                    brand)

    # ── Approval / sign-off ──
    approvals = job_aid.get("approvals") or []
    if approvals:
        _heading(doc, "Approval")
        _grid_table(doc, ["Role", "Name", "Signature", "Date"],
                    [[a.get("role", ""), a.get("name", ""), "", a.get("date", "")] for a in approvals],
                    brand)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
