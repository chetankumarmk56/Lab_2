"""Lab 4 — Citizen Service Job Aid Generator API.

Inputs: a tested workflow (paste / upload / link) + a document type + a template
(approved library / upload .docx / link). Output: a branded .docx (returned
base64 so it downloads in one round-trip) plus a structured preview.
"""
import asyncio
import base64
import io
import json
import os
import re
import tempfile
from pathlib import Path
from typing import Optional

import httpx
from docx import Document
from fastapi import APIRouter, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse

from ..agents.lab4_job_aid import generate_job_aid
from ..config import DATA_DIR
from ..lab4_docx import render_job_aid
from ..lab4_templates import (
    default_template_id,
    list_templates,
    template_meta,
    template_path,
)

router = APIRouter(prefix="/api/lab4", tags=["Lab 4 — Job Aid"])

LAB4_DATA = DATA_DIR / "lab4"
DOC_TYPES = {"Job Aid", "User Manual", "Training Guide", "Training"}


def _extract_text(filename: str, raw: bytes) -> str:
    """Pull plain text from an uploaded .docx / .txt / .md file."""
    if (filename or "").lower().endswith(".docx"):
        doc = Document(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs)
    return raw.decode("utf-8-sig", errors="replace")


def _fetch_text_sync(url: str) -> str:
    r = httpx.get(url, timeout=15, follow_redirects=True)
    r.raise_for_status()
    if "text/html" in r.headers.get("content-type", ""):
        html = re.sub(r"<(script|style)[\s\S]*?</\1>", " ", r.text, flags=re.I)
        return re.sub(r"[ \t]+", " ", re.sub(r"<[^>]+>", " ", html)).strip()
    return r.text


def _fetch_bytes_sync(url: str) -> bytes:
    r = httpx.get(url, timeout=15, follow_redirects=True)
    r.raise_for_status()
    return r.content


def _parse_job_aid(text: str) -> Optional[dict]:
    if not text:
        return None
    i, j = text.find("{"), text.rfind("}")
    if i == -1 or j == -1:
        return None
    try:
        obj = json.loads(text[i : j + 1])
        return obj if isinstance(obj, dict) else None
    except json.JSONDecodeError:
        return None


def _slug(text: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", (text or "job-aid").lower()).strip("-")
    return s[:60] or "job-aid"


@router.get("/templates")
async def templates():
    return {"templates": list_templates()}


@router.get("/sample")
async def sample():
    return FileResponse(
        LAB4_DATA / "sample_workflow.md", media_type="text/markdown", filename="sample_workflow.md"
    )


@router.post("/generate")
async def generate(
    doc_type: str = Form("Job Aid"),
    workflow_text: str = Form(""),
    workflow_url: str = Form(""),
    template_id: str = Form(""),
    template_url: str = Form(""),
    workflow_file: Optional[UploadFile] = File(None),
    template_file: Optional[UploadFile] = File(None),
):
    if doc_type not in DOC_TYPES:
        doc_type = "Job Aid"

    # ---- resolve the workflow text (upload > link > paste) ----
    text = ""
    if workflow_file is not None:
        text = _extract_text(workflow_file.filename, await workflow_file.read())
    elif workflow_url.strip():
        try:
            text = await asyncio.to_thread(_fetch_text_sync, workflow_url.strip())
        except Exception as exc:  # noqa: BLE001
            raise HTTPException(400, f"Couldn't fetch the workflow link: {exc}")
    elif workflow_text.strip():
        text = workflow_text
    text = text.strip()
    if len(text) < 20:
        raise HTTPException(
            400, "Provide the tested workflow — paste the steps, upload a file, or give a link."
        )

    # ---- resolve the template canvas (upload > link > library) ----
    tmp_template: Optional[str] = None
    try:
        if template_file is not None and (template_file.filename or "").lower().endswith(".docx"):
            raw = await template_file.read()
            fd, tmp_template = tempfile.mkstemp(suffix=".docx")
            os.close(fd)
            Path(tmp_template).write_bytes(raw)
            try:
                Document(tmp_template)
            except Exception:
                raise HTTPException(400, "That uploaded template isn't a valid .docx file.")
            base_path, template_name, agency = Path(tmp_template), template_file.filename, "the agency"
        elif template_url.strip():
            try:
                raw = await asyncio.to_thread(_fetch_bytes_sync, template_url.strip())
                fd, tmp_template = tempfile.mkstemp(suffix=".docx")
                os.close(fd)
                Path(tmp_template).write_bytes(raw)
                Document(tmp_template)
            except HTTPException:
                raise
            except Exception as exc:  # noqa: BLE001
                raise HTTPException(400, f"Couldn't use the template link as a .docx: {exc}")
            base_path, template_name, agency = Path(tmp_template), template_url.strip(), "the agency"
        else:
            tid = template_id or default_template_id()
            base_path = template_path(tid) or template_path(default_template_id())
            template_name, agency = template_meta(tid)

        # ---- run the agent, parse, render ----
        result = await generate_job_aid(text, doc_type, agency)
        if result["error"] and not result["result"]:
            raise HTTPException(502, f"Agent error: {result['error']}")
        job_aid = _parse_job_aid(result["result"])
        if job_aid is None:
            raise HTTPException(502, "The agent did not return a valid job aid. Please try again.")
        job_aid.setdefault("document_type", doc_type)

        docx_bytes = render_job_aid(job_aid, base_path)
    finally:
        if tmp_template:
            try:
                os.unlink(tmp_template)
            except OSError:
                pass

    return {
        "job_aid": job_aid,
        "template_name": template_name,
        "filename": f"{_slug(job_aid.get('title'))}.docx",
        "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
    }
