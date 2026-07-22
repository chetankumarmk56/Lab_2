"""Lab 4 — approved template library.

Each template is a branded, empty .docx "canvas" (agency name in the header,
brand-coloured Title/Heading styles, a footer). The job-aid renderer opens the
chosen canvas and appends content, so the output inherits the template's
branding exactly — the faithful way to "match the template".

The .docx canvases are generated on first use into data/lab4/templates/ (so no
build step is required), and the SAME open-and-append path is reused when a user
uploads their own .docx template.
"""
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor

from .config import DATA_DIR

TEMPLATES_DIR = DATA_DIR / "lab4" / "templates"

TEMPLATES = [
    {
        "id": "state-dmv",
        "name": "State DMV — Official Blue",
        "agency": "State Department of Motor Vehicles",
        "brand": "1F3A93",
        "body_font": "Calibri",
        "heading_font": "Calibri",
        "footer": "Official Job Aid · For Counter Staff Use Only",
    },
    {
        "id": "county-services",
        "name": "County Citizen Services — Green",
        "agency": "County Citizen Services Office",
        "brand": "1F7A44",
        "body_font": "Georgia",
        "heading_font": "Georgia",
        "footer": "County Citizen Services · Internal Use",
    },
    {
        "id": "training-division",
        "name": "Statewide Training — Slate",
        "agency": "Statewide Training Division",
        "brand": "334155",
        "body_font": "Arial",
        "heading_font": "Arial",
        "footer": "Training Division · Controlled Document",
    },
]


def _spec(template_id: str) -> Optional[dict]:
    return next((t for t in TEMPLATES if t["id"] == template_id), None)


def _style_font(style, name=None, size=None, color=None, bold=None) -> None:
    f = style.font
    if name:
        f.name = name
    if size is not None:
        f.size = Pt(size)
    if color:
        f.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        f.bold = bold


def _build(spec: dict) -> Document:
    doc = Document()
    _style_font(doc.styles["Normal"], name=spec["body_font"], size=11)
    _style_font(doc.styles["Title"], name=spec["heading_font"], size=28, color=spec["brand"], bold=True)
    _style_font(doc.styles["Heading 1"], name=spec["heading_font"], size=15, color=spec["brand"], bold=True)
    _style_font(doc.styles["Heading 2"], name=spec["heading_font"], size=12, color=spec["brand"], bold=True)

    sec = doc.sections[0]
    head = sec.header.paragraphs[0]
    head.text = spec["agency"]
    if head.runs:
        r = head.runs[0]
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor.from_string(spec["brand"])

    foot = sec.footer.paragraphs[0]
    foot.text = spec["footer"]
    if foot.runs:
        r = foot.runs[0]
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string("9AA0AA")

    # Leave the body empty — it's a canvas the renderer appends into.
    for p in list(doc.paragraphs):
        p._element.getparent().remove(p._element)
    return doc


def ensure_templates() -> None:
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    for spec in TEMPLATES:
        path = TEMPLATES_DIR / f"{spec['id']}.docx"
        if not path.exists():
            _build(spec).save(str(path))


def list_templates() -> list[dict]:
    ensure_templates()
    return [
        {"id": t["id"], "name": t["name"], "agency": t["agency"], "brand": "#" + t["brand"]}
        for t in TEMPLATES
    ]


def default_template_id() -> str:
    return TEMPLATES[0]["id"]


def template_path(template_id: str) -> Optional[Path]:
    ensure_templates()
    spec = _spec(template_id)
    if not spec:
        return None
    return TEMPLATES_DIR / f"{spec['id']}.docx"


def template_meta(template_id: str) -> tuple[str, str]:
    """Return (display_name, agency) for a library template id."""
    spec = _spec(template_id)
    if not spec:
        return template_id, "the agency"
    return spec["name"], spec["agency"]


def template_brand(template_id: str) -> str:
    """Brand hex (no leading #) for a library template; blueprint blue by default."""
    spec = _spec(template_id)
    return spec["brand"] if spec else "1F3A93"
